"""
Resume Parser Service
Extracts text content from PDF and DOCX resume files.
"""

import io
import re
from typing import Optional

from PyPDF2 import PdfReader
from docx import Document


class ResumeParser:
    """Parses resume files and extracts structured text content."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

    def __init__(self):
        pass

    def parse(self, file_content: bytes, filename: str) -> dict:
        """
        Parse a resume file and return extracted text and metadata.

        Args:
            file_content: Raw bytes of the uploaded file.
            filename: Original filename (used to determine format).

        Returns:
            Dictionary with 'text', 'filename', 'format', and 'sections'.
        """
        ext = self._get_extension(filename)

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        if ext == ".pdf":
            raw_text = self._parse_pdf(file_content)
        elif ext == ".docx":
            raw_text = self._parse_docx(file_content)
        else:
            raise ValueError(f"Unsupported format: {ext}")

        sections = self._extract_sections(raw_text)

        return {
            "text": raw_text,
            "filename": filename,
            "format": ext.lstrip(".").upper(),
            "sections": sections,
            "line_count": len(raw_text.strip().splitlines()),
            "word_count": len(raw_text.split()),
            "char_count": len(raw_text),
        }

    def _get_extension(self, filename: str) -> str:
        """Extract file extension in lowercase."""
        if "." not in filename:
            return ""
        return "." + filename.rsplit(".", 1)[-1].lower()

    def _parse_pdf(self, file_content: bytes) -> str:
        """Extract text from a PDF file."""
        try:
            reader = PdfReader(io.BytesIO(file_content))
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    def _parse_docx(self, file_content: bytes) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    def _extract_sections(self, text: str) -> dict:
        """
        Identify common resume sections from the text.

        Returns a dict mapping section names to their content.
        """
        section_patterns = {
            "contact": r"(?i)(contact\s*info(?:rmation)?|personal\s*(?:info(?:rmation)?|details))",
            "summary": r"(?i)(summary|objective|professional\s*summary|career\s*summary|profile|about\s*me)",
            "experience": r"(?i)(work\s*experience|professional\s*experience|experience|employment\s*history|work\s*history)",
            "education": r"(?i)(education|academic\s*(?:background|qualifications)|qualifications)",
            "skills": r"(?i)(skills|technical\s*skills|core\s*competencies|competencies|proficiencies|areas\s*of\s*expertise)",
            "certifications": r"(?i)(certifications?|licenses?|credentials|professional\s*development)",
            "projects": r"(?i)(projects|key\s*projects|notable\s*projects|personal\s*projects)",
            "awards": r"(?i)(awards?|honors?|achievements?|accomplishments?)",
            "publications": r"(?i)(publications?|research|papers)",
            "languages": r"(?i)(languages?|language\s*proficiency)",
            "volunteer": r"(?i)(volunteer(?:ing)?|community\s*(?:service|involvement))",
            "references": r"(?i)(references?)",
        }

        lines = text.split("\n")
        found_sections = {}
        current_section = "header"
        current_content = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            matched_section = None
            # Only consider short lines as potential section headers (max ~6 words)
            if len(stripped.split()) <= 6:
                for section_name, pattern in section_patterns.items():
                    # Match full line as a section header
                    full_line_pattern = (
                        r"(?i)^[\s\-\*]*"
                        + pattern.replace("(?i)", "")
                        + r"[\s:\-]*$"
                    )
                    if re.match(full_line_pattern, stripped):
                        matched_section = section_name
                        break

            if matched_section:
                # Save previous section
                if current_content:
                    found_sections[current_section] = "\n".join(current_content)
                current_section = matched_section
                current_content = []
            else:
                current_content.append(stripped)

        # Save last section
        if current_content:
            found_sections[current_section] = "\n".join(current_content)

        return found_sections
